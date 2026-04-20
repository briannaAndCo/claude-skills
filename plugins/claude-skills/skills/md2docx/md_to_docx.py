#!/usr/bin/env python3
"""Convert markdown files to professionally styled .docx documents."""

import sys
import re
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

# Color palette
HEADING_COLOR = RGBColor(0x1B, 0x3A, 0x5C)  # Dark navy blue
BODY_COLOR = RGBColor(0x33, 0x33, 0x33)  # Dark gray
SUBTITLE_COLOR = RGBColor(0x6B, 0x70, 0x7B)  # Medium gray
CODE_TEXT_COLOR = RGBColor(0x2D, 0x2D, 0x2D)
CODE_BG = "EAECF0"
INLINE_CODE_BG = "EAECF0"
TABLE_HEADER_BG = "E8EDF3"
TABLE_HEADER_TEXT = RGBColor(0x1B, 0x3A, 0x5C)
TABLE_ALT_ROW_BG = "F7F8FA"
ACCENT_COLOR = "1B3A5C"
HR_COLOR = "CBD0D8"

CODE_FONT = "Consolas"
BODY_FONT = "Calibri"
HEADING_FONT = "Calibri"


def set_cell_shading(cell, color):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def set_cell_margins(cell, top=80, bottom=80, left=120, right=120):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = parse_xml(
        f'<w:tcMar {nsdecls("w")}>'
        f'<w:top w:w="{top}" w:type="dxa"/>'
        f'<w:bottom w:w="{bottom}" w:type="dxa"/>'
        f'<w:start w:w="{left}" w:type="dxa"/>'
        f'<w:end w:w="{right}" w:type="dxa"/>'
        f'</w:tcMar>'
    )
    tcMar_old = tcPr.find(qn('w:tcMar'))
    if tcMar_old is not None:
        tcPr.remove(tcMar_old)
    tcPr.append(tcMar)


def remove_table_borders(table):
    """Remove all borders from a table."""
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        '<w:top w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '<w:left w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '<w:bottom w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '<w:right w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '<w:insideH w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '<w:insideV w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '</w:tblBorders>'
    )
    old_borders = tblPr.find(qn('w:tblBorders'))
    if old_borders is not None:
        tblPr.remove(old_borders)
    tblPr.append(borders)


def set_table_borders(table, color="CBD0D8"):
    """Set clean borders on a data table."""
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'<w:top w:val="single" w:sz="4" w:space="0" w:color="{color}"/>'
        f'<w:left w:val="single" w:sz="4" w:space="0" w:color="{color}"/>'
        f'<w:bottom w:val="single" w:sz="4" w:space="0" w:color="{color}"/>'
        f'<w:right w:val="single" w:sz="4" w:space="0" w:color="{color}"/>'
        f'<w:insideH w:val="single" w:sz="4" w:space="0" w:color="{color}"/>'
        f'<w:insideV w:val="single" w:sz="4" w:space="0" w:color="{color}"/>'
        f'</w:tblBorders>'
    )
    old_borders = tblPr.find(qn('w:tblBorders'))
    if old_borders is not None:
        tblPr.remove(old_borders)
    tblPr.append(borders)


def add_code_block(doc, lines):
    """Add a code block using a single-cell table with background shading."""
    code_text = "\n".join(lines)

    # Create a 1x1 table for the code block
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    # Set full width
    tbl = table._tbl
    tblPr = tbl.tblPr
    tblW = parse_xml(f'<w:tblW {nsdecls("w")} w:type="pct" w:w="5000"/>')
    old_tblW = tblPr.find(qn('w:tblW'))
    if old_tblW is not None:
        tblPr.remove(old_tblW)
    tblPr.append(tblW)

    remove_table_borders(table)

    cell = table.rows[0].cells[0]
    set_cell_shading(cell, CODE_BG)
    set_cell_margins(cell, top=160, bottom=160, left=220, right=220)

    # Clear default paragraph
    cell.text = ''
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = Pt(18)

    run = p.add_run(code_text)
    run.font.name = CODE_FONT
    run.font.size = Pt(10)
    run.font.color.rgb = CODE_TEXT_COLOR

    # Add spacing after code block
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_before = Pt(2)
    spacer.paragraph_format.space_after = Pt(2)
    sp_fmt = spacer.paragraph_format
    sp_fmt.line_spacing = Pt(4)


def add_inline_formatted(paragraph, text, base_size=11, base_color=None):
    """Add text with inline code and bold formatting."""
    if base_color is None:
        base_color = BODY_COLOR

    parts = re.split(r'(`[^`]+`)', text)
    for part in parts:
        if part.startswith('`') and part.endswith('`'):
            code_text = part[1:-1]
            run = paragraph.add_run(code_text)
            run.font.name = CODE_FONT
            run.font.size = Pt(base_size)
            run.font.color.rgb = RGBColor(0x1B, 0x3A, 0x5C)
        else:
            bold_parts = re.split(r'(\*\*[^*]+\*\*)', part)
            for bp in bold_parts:
                if bp.startswith('**') and bp.endswith('**'):
                    run = paragraph.add_run(bp[2:-2])
                    run.bold = True
                    run.font.size = Pt(base_size)
                    run.font.name = BODY_FONT
                    run.font.color.rgb = base_color
                else:
                    cleaned = re.sub(r'\[([^\]]+)\]\([^)]+\)', r'\1', bp)
                    if cleaned:
                        run = paragraph.add_run(cleaned)
                        run.font.size = Pt(base_size)
                        run.font.name = BODY_FONT
                        run.font.color.rgb = base_color


def convert_md_to_docx(md_path, docx_path):
    doc = Document()

    # Set margins
    for section in doc.sections:
        section.top_margin = Inches(1.25)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.25)
        section.right_margin = Inches(1.25)

    # Configure default style
    style = doc.styles['Normal']
    style.font.name = BODY_FONT
    style.font.size = Pt(11)
    style.font.color.rgb = BODY_COLOR
    style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    style.paragraph_format.line_spacing = 1.15
    style.paragraph_format.space_after = Pt(8)
    style.paragraph_format.space_before = Pt(2)

    # Configure heading styles
    for level in range(1, 4):
        hs = doc.styles[f'Heading {level}']
        hs.font.name = HEADING_FONT
        hs.font.color.rgb = HEADING_COLOR
        hs.font.bold = True
        if level == 1:
            hs.font.size = Pt(26)
            hs.paragraph_format.space_before = Pt(0)
            hs.paragraph_format.space_after = Pt(14)
        elif level == 2:
            hs.font.size = Pt(18)
            hs.paragraph_format.space_before = Pt(28)
            hs.paragraph_format.space_after = Pt(12)
        else:
            hs.font.size = Pt(14)
            hs.paragraph_format.space_before = Pt(20)
            hs.paragraph_format.space_after = Pt(10)

    # Configure list styles
    for list_style_name in ['List Bullet', 'List Number']:
        try:
            ls = doc.styles[list_style_name]
            ls.font.name = BODY_FONT
            ls.font.size = Pt(11)
            ls.font.color.rgb = BODY_COLOR
            ls.paragraph_format.space_after = Pt(3)
            ls.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
            ls.paragraph_format.line_spacing = 1.15
        except KeyError:
            pass

    with open(md_path, 'r') as f:
        lines = f.readlines()

    i = 0
    in_code_block = False
    code_lines = []
    in_table = False
    table_rows = []
    is_first_heading = True

    while i < len(lines):
        line = lines[i].rstrip('\n')

        # Code blocks
        if line.startswith('```'):
            if in_code_block:
                add_code_block(doc, code_lines)
                code_lines = []
                in_code_block = False
            else:
                in_code_block = True
            i += 1
            continue

        if in_code_block:
            code_lines.append(line)
            i += 1
            continue

        # Tables
        if line.startswith('|') and line.endswith('|'):
            if not in_table:
                in_table = True
                table_rows = []
            if not re.match(r'^\|[\s\-:|]+\|$', line):
                table_rows.append(line)
            i += 1
            continue
        elif in_table:
            _render_table(doc, table_rows)
            in_table = False
            table_rows = []
            continue

        # H1
        if line.startswith('# '):
            if is_first_heading:
                _add_title_block(doc, line[2:])
                is_first_heading = False
            else:
                h = doc.add_heading(level=1)
                h.clear()
                add_inline_formatted(h, line[2:], base_size=26, base_color=HEADING_COLOR)
            i += 1
            continue

        # H2
        if line.startswith('## '):
            h = doc.add_heading(level=2)
            h.clear()
            add_inline_formatted(h, line[3:], base_size=18, base_color=HEADING_COLOR)
            i += 1
            continue

        # H3
        if line.startswith('### '):
            h = doc.add_heading(level=3)
            h.clear()
            add_inline_formatted(h, line[4:], base_size=14, base_color=HEADING_COLOR)
            i += 1
            continue

        # Horizontal rule
        if line.startswith('---'):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(20)
            p.paragraph_format.space_after = Pt(20)
            pPr = p._p.get_or_add_pPr()
            pBdr = parse_xml(
                f'<w:pBdr {nsdecls("w")}>'
                f'<w:bottom w:val="single" w:sz="4" w:space="4" w:color="{HR_COLOR}"/>'
                f'</w:pBdr>'
            )
            pPr.append(pBdr)
            i += 1
            continue

        # Blockquote
        if line.startswith('> '):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(0.8)
            pPr = p._p.get_or_add_pPr()
            pBdr = parse_xml(
                f'<w:pBdr {nsdecls("w")}>'
                f'<w:left w:val="single" w:sz="18" w:space="8" w:color="{ACCENT_COLOR}"/>'
                f'</w:pBdr>'
            )
            pPr.append(pBdr)
            run = p.add_run(line[2:])
            run.font.size = Pt(11)
            run.font.name = BODY_FONT
            run.font.color.rgb = SUBTITLE_COLOR
            run.italic = True
            i += 1
            continue

        # Unordered list
        if re.match(r'^[-*] ', line):
            item_text = re.sub(r'^[-*] ', '', line)
            p = doc.add_paragraph(style='List Bullet')
            add_inline_formatted(p, item_text)
            i += 1
            continue

        # Ordered list
        if re.match(r'^\d+\. ', line):
            item_text = re.sub(r'^\d+\. ', '', line)
            p = doc.add_paragraph(style='List Number')
            add_inline_formatted(p, item_text)
            i += 1
            continue

        # Empty line
        if line.strip() == '':
            i += 1
            continue

        # Regular paragraph
        p = doc.add_paragraph()
        add_inline_formatted(p, line)
        i += 1

    # Handle remaining table
    if in_table and table_rows:
        _render_table(doc, table_rows)

    doc.save(docx_path)
    print(f"Saved: {docx_path}")


def _add_title_block(doc, title_text):
    """Add a styled title with an accent bar above it."""
    bar_table = doc.add_table(rows=1, cols=1)
    bar_table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl = bar_table._tbl
    tblPr = tbl.tblPr
    tblW = parse_xml(f'<w:tblW {nsdecls("w")} w:type="pct" w:w="5000"/>')
    old_tblW = tblPr.find(qn('w:tblW'))
    if old_tblW is not None:
        tblPr.remove(old_tblW)
    tblPr.append(tblW)
    remove_table_borders(bar_table)

    cell = bar_table.rows[0].cells[0]
    set_cell_shading(cell, ACCENT_COLOR)
    set_cell_margins(cell, top=0, bottom=0, left=0, right=0)
    tr = bar_table.rows[0]._tr
    trPr = tr.get_or_add_trPr()
    trHeight = parse_xml(f'<w:trHeight {nsdecls("w")} w:val="80" w:hRule="exact"/>')
    trPr.append(trHeight)
    cell.text = ''
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run('')
    run.font.size = Pt(2)

    # Spacer
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_before = Pt(6)
    spacer.paragraph_format.space_after = Pt(0)
    sf = spacer.paragraph_format
    sf.line_spacing = Pt(6)

    # Title heading
    h = doc.add_heading(level=1)
    h.clear()
    run = h.add_run(title_text)
    run.font.name = HEADING_FONT
    run.font.size = Pt(28)
    run.font.color.rgb = HEADING_COLOR
    run.bold = True
    h.paragraph_format.space_before = Pt(0)
    h.paragraph_format.space_after = Pt(4)


def _render_table(doc, table_rows):
    """Render a markdown table as a styled Word table."""
    if not table_rows:
        return

    num_cols = len([c for c in table_rows[0].split('|') if c.strip()])
    table = doc.add_table(rows=len(table_rows), cols=num_cols)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    # Full width
    tbl = table._tbl
    tblPr = tbl.tblPr
    tblW = parse_xml(f'<w:tblW {nsdecls("w")} w:type="pct" w:w="5000"/>')
    old_tblW = tblPr.find(qn('w:tblW'))
    if old_tblW is not None:
        tblPr.remove(old_tblW)
    tblPr.append(tblW)

    set_table_borders(table, HR_COLOR)

    for row_idx, row_text in enumerate(table_rows):
        cells_text = [c.strip() for c in row_text.split('|') if c.strip() != '']
        for col_idx, cell_text in enumerate(cells_text):
            if col_idx < num_cols:
                cell = table.rows[row_idx].cells[col_idx]
                set_cell_margins(cell, top=80, bottom=80, left=120, right=120)
                cell.text = ''
                p = cell.paragraphs[0]
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(0)

                clean = re.sub(r'`([^`]+)`', r'\1', cell_text)
                clean = re.sub(r'\*\*([^*]+)\*\*', r'\1', clean)
                run = p.add_run(clean)
                run.font.size = Pt(10)
                run.font.name = BODY_FONT

                if row_idx == 0:
                    run.bold = True
                    run.font.color.rgb = TABLE_HEADER_TEXT
                    set_cell_shading(cell, TABLE_HEADER_BG)
                elif row_idx % 2 == 0:
                    run.font.color.rgb = BODY_COLOR
                    set_cell_shading(cell, TABLE_ALT_ROW_BG)
                else:
                    run.font.color.rgb = BODY_COLOR

    # Spacing after table
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_before = Pt(4)
    spacer.paragraph_format.space_after = Pt(4)
    sf = spacer.paragraph_format
    sf.line_spacing = Pt(6)


if __name__ == '__main__':
    if len(sys.argv) < 3:
        print("Usage: md_to_docx.py <input.md> <output.docx>")
        sys.exit(1)
    convert_md_to_docx(sys.argv[1], sys.argv[2])
